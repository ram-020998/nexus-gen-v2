"""
Word Document Service - Generate Word documents for design exports
"""
from config import Config
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordService:
    """Handle Word document generation"""

    def __init__(self):
        self.config = Config

    def create_design_document(self, request_id: int, design_data: dict) -> str:
        """Create Word document for design data"""
        # Create output directory
        output_dir = self.config.OUTPUT_FOLDER / str(request_id)
        output_dir.mkdir(parents=True, exist_ok=True)

        # Create Word document
        doc = Document()

        # Add title
        title = doc.add_heading('Design Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add overview section
        if 'design_document' in design_data and 'overview' in design_data['design_document']:
            doc.add_heading('Overview', level=1)
            doc.add_paragraph(design_data['design_document']['overview'])

        # Add existing objects to modify
        if 'design_document' in design_data and 'existing_objects_to_modify' in design_data['design_document']:
            existing_objects = design_data['design_document']['existing_objects_to_modify']
            if existing_objects:
                doc.add_heading('Objects to Modify', level=1)

                for obj in existing_objects:
                    doc.add_heading(obj.get('name', 'Unknown Object'), level=2)
                    doc.add_paragraph(f"Type: {obj.get('type', 'N/A')}")
                    doc.add_paragraph(f"Current Description: {obj.get('current_description', 'No description')}")
                    doc.add_paragraph(f"Proposed Changes: {obj.get('proposed_changes', 'No changes specified')}")

                    if obj.get('new_methods'):
                        doc.add_paragraph("New Methods:")
                        for method in obj['new_methods']:
                            p = doc.add_paragraph(f"• {method}")
                            p.style = 'List Bullet'

        # Add new objects section
        if 'design_document' in design_data and 'new_objects' in design_data['design_document']:
            new_objects = design_data['design_document']['new_objects']
            if new_objects:
                doc.add_heading('New Objects', level=1)

                for obj in new_objects:
                    doc.add_heading(obj.get('name', 'Unknown Object'), level=2)
                    doc.add_paragraph(f"Type: {obj.get('type', 'N/A')}")
                    doc.add_paragraph(f"Description: {obj.get('description', 'No description')}")

                    if obj.get('methods'):
                        doc.add_paragraph("Methods:")
                        for method in obj['methods']:
                            p = doc.add_paragraph(f"• {method}")
                            p.style = 'List Bullet'

        # Add implementation notes
        if 'design_document' in design_data and 'implementation_notes' in design_data['design_document']:
            doc.add_heading('Implementation Notes', level=1)
            for note in design_data['design_document']['implementation_notes']:
                p = doc.add_paragraph(f"• {note}")
                p.style = 'List Bullet'

        # Add dependencies
        if 'design_document' in design_data and 'dependencies' in design_data['design_document']:
            doc.add_heading('Dependencies', level=1)
            for dep in design_data['design_document']['dependencies']:
                p = doc.add_paragraph(f"• {dep}")
                p.style = 'List Bullet'

        # Save document
        filename = f"design_document_{request_id}.docx"
        file_path = output_dir / filename
        doc.save(str(file_path))

        return str(file_path)
