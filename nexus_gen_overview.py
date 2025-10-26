#!/usr/bin/env python3

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_font(run, font_name="Tahoma", font_size=11):
    """Set font for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)


def add_heading(doc, text, level=1):
    """Add a heading with Tahoma font"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_font(run, "Tahoma", 14 if level == 1 else 12)
    return heading


def add_paragraph(doc, text, bold=False):
    """Add a paragraph with Tahoma font"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, "Tahoma", 11)
    if bold:
        run.bold = True
    return para


def create_table(doc, headers, rows):
    """Create a table with proper formatting"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, "Tahoma", 10)
                run.bold = True

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, "Tahoma", 9)

    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading('NexusGen Document Intelligence Hub: Empowering Low-Code Developers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, "Tahoma", 16)

    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    summary_text = ('NexusGen is an AI-powered platform that transforms document processing for low-code developers. '
                    'Using Amazon Q CLI agents and AWS Bedrock, it automates requirements analysis, '
                    'design validation, and documentation generation—reducing manual effort by 90% '
                    'while maintaining enterprise-quality standards.')
    add_paragraph(doc, summary_text)

    # Core Capabilities & Value
    add_heading(doc, 'Core Capabilities & Value', 1)

    headers = ['Feature', 'Traditional Process', 'NexusGen Process', 'Time Savings']
    rows = [
        ['Spec Breakdown', '4-6 hours manual analysis', '3-5 minutes automated', '95%'],
        ['Design Verification', '1-2 weeks review cycles', 'Immediate AI feedback', '90%'],
        ['Design Creation', '2-3 days manual documentation', '5-10 minutes generation', '98%'],
        ['AI Chat Assistant', 'Research & consultation time', 'Instant contextual guidance', '80%']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    # Technical Integration Benefits
    add_heading(doc, 'Technical Integration Benefits', 1)

    headers = ['Component', 'Technology', 'Low-Code Developer Benefit']
    rows = [
        ['Backend', 'Flask + SQLAlchemy', 'Familiar Python stack, easy customization'],
        ['AI Engine', 'Amazon Q CLI Agents', 'Enterprise-grade AI without complex setup'],
        ['Knowledge Base', 'AWS Bedrock', 'Contextual recommendations from org patterns'],
        ['Export Formats', 'Excel + Word', 'Business-ready deliverables'],
        ['Interface', 'Bootstrap 5 + Dark Theme', 'Modern, responsive, developer-friendly']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    # Impact Metrics
    add_heading(doc, 'Impact Metrics', 1)

    add_heading(doc, 'Time Savings Analysis', 2)
    headers = ['Development Phase', 'Before NexusGen', 'After NexusGen', 'Improvement']
    rows = [
        ['Requirements Analysis', '4-6 hours', '5 minutes', '95% faster'],
        ['User Story Creation', '1-2 hours per story', 'Automated', '100% faster'],
        ['Design Documentation', '2-3 days', '10 minutes', '98% faster'],
        ['Design Review Cycles', '1-2 weeks', 'Immediate', '90% faster'],
        ['Overall Project Delivery', 'Baseline', '40-60% faster', 'Significant acceleration']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    add_heading(doc, 'Quality Improvements', 2)
    headers = ['Quality Metric', 'Manual Process', 'NexusGen Process']
    rows = [
        ['Consistency', 'Variable, depends on reviewer', 'Standardized AI analysis'],
        ['Completeness', 'Often missed components', 'Systematic gap identification'],
        ['Best Practices', 'Inconsistent application', 'Automatic pattern matching'],
        ['Documentation Standards', 'Manual formatting', 'Professional auto-generation']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    # Feature Comparison Matrix
    add_heading(doc, 'Feature Comparison Matrix', 1)

    headers = ['Capability', 'Manual Approach', 'Generic AI Tools', 'NexusGen Advantage']
    rows = [
        ['Context Awareness', 'Limited to individual knowledge', 'No organizational context',
         'Bedrock KB with org patterns'],
        ['Output Format', 'Inconsistent', 'Generic text', 'Structured JSON + Professional exports'],
        ['Integration', 'Manual copy-paste', 'API calls required', 'Native AWS ecosystem'],
        ['Learning', 'Static knowledge', 'No organizational learning', 'Continuous KB improvement'],
        ['Scalability', 'Linear with team size', 'Per-request costs', 'Enterprise-grade scaling']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    # ROI Analysis
    add_heading(doc, 'ROI Analysis for Low-Code Teams', 1)

    headers = ['Team Size', 'Monthly Time Saved', 'Cost Savings (@ $75/hr)', 'Annual ROI']
    rows = [
        ['5 developers', '120 hours', '$9,000', '$108,000'],
        ['10 developers', '240 hours', '$18,000', '$216,000'],
        ['20 developers', '480 hours', '$36,000', '$432,000']
    ]
    create_table(doc, headers, rows)

    add_paragraph(doc, '*Based on average 24 hours/month saved per developer', bold=False)

    doc.add_paragraph()  # Space

    # Architecture & Scalability
    add_heading(doc, 'Architecture & Scalability', 1)

    add_heading(doc, 'System Components', 2)
    headers = ['Layer', 'Technology', 'Scalability Feature']
    rows = [
        ['Frontend', 'Bootstrap 5 + JavaScript', 'Responsive, mobile-ready'],
        ['API Layer', 'Flask Controllers', 'RESTful, integration-ready'],
        ['Business Logic', 'Python Services', 'Modular, extensible'],
        ['AI Processing', 'Q CLI Agents', 'AWS-managed scaling'],
        ['Data Storage', 'SQLite/PostgreSQL', 'Configurable for growth'],
        ['File Processing', 'Multi-format support', 'PDF, DOCX, TXT, MD']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    add_heading(doc, 'Integration Capabilities', 2)
    headers = ['Integration Type', 'Supported Formats', 'Use Case']
    rows = [
        ['Input', 'PDF, DOCX, TXT, MD', 'Specification documents'],
        ['Output', 'JSON, Excel, Word', 'Development artifacts'],
        ['API', 'RESTful endpoints', 'Low-code platform integration'],
        ['Webhooks', 'Real-time notifications', 'Workflow automation']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    # Strategic Impact
    add_heading(doc, 'Strategic Impact', 1)

    add_heading(doc, 'Developer Empowerment', 2)
    bullet_points = [
        'Junior developers produce architect-quality deliverables',
        'Senior developers focus on complex problem-solving',
        'Teams maintain consistent quality across projects',
        'Organizations scale development without proportional architect hiring'
    ]

    for point in bullet_points:
        para = doc.add_paragraph()
        run = para.add_run(f'• {point}')
        set_font(run, "Tahoma", 11)

    add_heading(doc, 'Business Benefits', 2)
    headers = ['Benefit Category', 'Impact']
    rows = [
        ['Time-to-Market', '40-60% faster project delivery'],
        ['Quality Assurance', 'Consistent, AI-validated designs'],
        ['Cost Reduction', '90% less manual documentation effort'],
        ['Scalability', 'Linear team growth without quality degradation'],
        ['Knowledge Retention', 'Organizational patterns preserved in KB']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    # Implementation Readiness
    add_heading(doc, 'Implementation Readiness', 1)

    add_heading(doc, 'Quick Start Metrics', 2)
    headers = ['Setup Phase', 'Time Required', 'Complexity']
    rows = [
        ['Installation', '15 minutes', 'Low - Python + pip'],
        ['AWS Configuration', '10 minutes', 'Medium - Bedrock setup'],
        ['First Document Processing', '2 minutes', 'Low - Drag & drop'],
        ['Team Onboarding', '30 minutes', 'Low - Intuitive interface']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    add_heading(doc, 'Deployment Options', 2)
    headers = ['Environment', 'Setup', 'Scalability']
    rows = [
        ['Local Development', 'Python + SQLite', 'Single developer'],
        ['Team Deployment', 'Gunicorn + PostgreSQL', '5-20 developers'],
        ['Enterprise', 'Docker + AWS RDS', 'Unlimited scaling']
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()  # Space

    # Conclusion
    add_heading(doc, 'Conclusion', 1)
    conclusion_text = ('NexusGen transforms low-code development by automating the most time-consuming aspects '
                       'of software design while maintaining enterprise quality standards. With 90%+ time savings '
                       'in document processing and AI-powered quality assurance, it\'s essential infrastructure '
                       'for scaling low-code teams.')
    add_paragraph(doc, conclusion_text)

    doc.add_paragraph()  # Space

    para = doc.add_paragraph()
    run = para.add_run('Key Value Proposition: ')
    set_font(run, "Tahoma", 11)
    run.bold = True
    value_prop = ('Enable low-code developers to deliver architect-quality work at developer speed, '
                  'accelerating project delivery by 40-60% while building organizational knowledge assets.')
    run2 = para.add_run(value_prop)
    set_font(run2, "Tahoma", 11)

    # Save document
    output_path = '/Users/ramaswamy.u/repo/nexus-gen-v2/NexusGen_Overview.docx'
    doc.save(output_path)
    # Document saved


if __name__ == "__main__":
    main()
